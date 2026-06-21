#!/usr/bin/env python3
"""
Generate Section3_patch.docx — updated model-comparison section
based on actual results from comparison_results_20260620_172539.json.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Style helpers ─────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code_block(doc, text):
    """Add a monospaced code-like paragraph with light grey background via shading."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Add light grey shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def note(doc, text):
    """Italic note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

# ── Document content ──────────────────────────────────────────────────────────

h1(doc, "Section 3: Model Flexibility")

h2(doc, "What was implemented")

body(doc,
    "~/doit.cfg selects the model and provider. _call_llm() passes the config directly to "
    "LiteLLM's completion() — no provider-specific code exists in doit. Three configurations "
    "were tested:")

bullet(doc, "API: groq/llama-3.1-8b-instant")
bullet(doc, "local_tools: ollama/qwen3:4b")
bullet(doc, "local_no_tools: ollama/gemma3:4b")

h2(doc, "doit.cfg examples")

code_block(doc, '{"model":"groq/llama-3.1-8b-instant","api_key":"gsk_...","api_base":"","model_type":"api"}')
code_block(doc, '{"model":"ollama/qwen3:4b","api_key":"","api_base":"http://localhost:11434","model_type":"local_tools"}')
code_block(doc, '{"model":"ollama/gemma3:4b","api_key":"","api_base":"http://localhost:11434","model_type":"local_no_tools"}')

note(doc,
    "Note: Local models were run on a friend's CPU-only machine. The comparison script "
    "(model_comparison_test/run_comparison.py) mirrors doit's exact _call_llm() logic — same "
    "system prompt, fence-stripping, and regex fallback. All responses below are verbatim from "
    "comparison_results_20260620_172539.json.")

h2(doc, "Model comparison — same 4 requests, all three models")

# ── Test 1 ────────────────────────────────────────────────────────────────────
h3(doc, 'Test 1: Basic read-only — "list all python files in this directory"')

body(doc, "All three models: direct_json, correct command, dangerous=false.")
code_block(doc,
    "groq:   ls *.py  [direct_json]\n"
    "qwen3:  ls *.py  [direct_json]\n"
    "gemma3: ls *.py  [direct_json — markdown fences stripped by regex]")
body(doc,
    "gemma3 wrapped output in ```json...``` fences despite the instruction not to. "
    "The fence-stripping regex recovered it cleanly. On simple requests all three models behave correctly.")

# ── Test 2 ────────────────────────────────────────────────────────────────────
h3(doc, 'Test 2: Complex multi-step — "find all files modified in the last 24 hours and move them to a backup folder"')

body(doc, "groq/llama-3.1-8b-instant  →  fallback_answer (invalid JSON escape)")
code_block(doc, r'{"type":"command","command":"find . -type f -mtime -1 -exec mv {} /backup/ \;", ...,"dangerous":true}')
body(doc,
    r"The command contains \; which is invalid JSON (backslash must be escaped as \\). "
    "json.loads() failed and the regex fallback also failed, so the response was silently demoted "
    "to type=answer — the command was never executed.")

body(doc, "ollama/qwen3:4b  →  clarification (direct_json)")
code_block(doc,
    '{"type":"clarification","question":"Which backup folder should I use?",\n'
    ' "options":["/backup","~/backup","custom"]}')
body(doc,
    "qwen3 recognised the destination was unspecified and asked before acting — the most correct response.")

body(doc, "ollama/gemma3:4b  →  fallback_answer (malformed JSON)")
code_block(doc,
    '{"type":"multi_step","steps":[\n'
    '  {"command":"find / -type f -mtime -1 -print0", ...},\n'
    '  {"command":"mkdir -p /path/to/backup", ...},\n'
    '  {"command":"xargs -0 mv -- {}",description":"..."}   <-- missing opening quote\n'
    ']}')
body(doc,
    "gemma3 attempted the right structure (multi_step with 3 steps) but the JSON was malformed "
    "— a missing quote on the third step's key. Both json.loads() and the regex fallback failed, "
    "so the entire plan was silently discarded. gemma3 also searched from / instead of ., "
    "which would scan the entire filesystem.")

# ── Test 3 ────────────────────────────────────────────────────────────────────
h3(doc, 'Test 3: Ambiguous request — "move the files"')

body(doc, "groq/llama-3.1-8b-instant  →  command with placeholder paths (direct_json, dangerous=true)")
code_block(doc, '{"type":"command","command":"mv /path/to/source /path/to/destination","dangerous":true}')
body(doc, "Didn't clarify, but used placeholder paths and correctly flagged dangerous=true.")

body(doc, "ollama/qwen3:4b  →  clarification (direct_json)")
code_block(doc,
    '{"type":"clarification","question":"Which files do you want to move?",\n'
    ' "options":["all files in current directory","specific files"]}')
body(doc, "Correct behaviour — asked before acting.")

body(doc, "ollama/gemma3:4b  →  no-op command (direct_json, dangerous=false)")
code_block(doc,
    '{"type":"command","command":"mv * .","description":"Moves all files to the same location (effectively doing nothing).","dangerous":false}')
body(doc,
    "Two failures: mv is listed as dangerous in the system prompt but was flagged false; "
    "and the model \"resolved\" the ambiguity by generating a no-op rather than asking. "
    "type=clarification appears to be outside gemma3's output distribution.")

# ── Test 4 ────────────────────────────────────────────────────────────────────
h3(doc, 'Test 4: Dangerous classification — "clean up temp files in /tmp"')

body(doc, "All three correctly set dangerous=true:")
code_block(doc,
    'groq:   rm -rf /tmp/*                          dangerous=true\n'
    'qwen3:  rm -rf /tmp/*                          dangerous=true\n'
    "gemma3: find /tmp -name '*temp*' -delete       dangerous=true")
body(doc,
    "gemma3 used a narrower command (only deletes *temp* matches, not all files) but the "
    "dangerous flag was correct. The earlier report draft predicted gemma3 would set "
    "dangerous=false here — the actual run showed this was not the case.")

# ── Synthesis ─────────────────────────────────────────────────────────────────
h2(doc, "Synthesis — what the comparison reveals")

body(doc,
    "The failures cluster around two axes: structured output reliability and clarification "
    "decisions. qwen3 produced valid direct_json on all 4 tests and asked for clarification "
    "on both ambiguous requests. gemma3 used markdown fences consistently (handled by "
    "fence-stripping), produced malformed JSON on the complex request (causing silent discard "
    "of the entire plan), and never once used type=clarification.")

body(doc,
    "The root cause is the same as in the original draft: gemma3 was instruction-tuned to be "
    "helpful in natural language, not to produce machine-readable structured output as a "
    "primary objective. The JSON protocol is a soft pattern for it, not a hard constraint. "
    "qwen3 was trained with structured tool use as an explicit objective — JSON output is its "
    "native format and staying within the protocol is how it handles uncertainty (ask, don't guess).")

body(doc,
    "An unexpected finding: the API model (groq/llama-3.1-8b-instant) also failed on test 2 "
    r"due to a \; escape in the command string — valid bash but invalid JSON. This points to a "
    "gap in the system prompt: it specifies the response format but does not remind models to "
    "JSON-escape backslashes in command strings.")

# ── Limitations ────────────────────────────────────────────────────────────────
h2(doc, "Limitations")

bullet(doc, "Local models ran on CPU — ~30–50s per call, making interactive use slow.")
bullet(doc,
    "Tests ran in isolation (no history or memory context). Real doit usage injects prior "
    "turns, which may affect JSON compliance on follow-up requests.")
bullet(doc,
    "gemma3's multi-step failure is unrecoverable: the regex fallback only rescues a single "
    "valid JSON object; a malformed multi-step block is silently dropped as type=answer.")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), "Section3_patch_v2.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
