"""Patch v5: replace Section 11 content, save as v6."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = r'C:\Users\nadav\OneDrive\university\2nd_math_degree\deep_learning_methods_for_texts_and_sequences\assignments\LLM_Shell_Agents\Assignment3_Report_v5.docx'
DEST = r'C:\Users\nadav\OneDrive\university\2nd_math_degree\deep_learning_methods_for_texts_and_sequences\assignments\LLM_Shell_Agents\Assignment3_Report_v6.docx'

doc = Document(SRC)
paras = doc.paragraphs

# Find Section 11 heading and the next H1 after it
sec11_idx = None
next_h1_idx = None
for i, p in enumerate(paras):
    text = p.text.strip()
    if sec11_idx is None and 'Section 11' in text and p.style.name.startswith('Heading 1'):
        sec11_idx = i
    elif sec11_idx is not None and p.style.name.startswith('Heading 1'):
        next_h1_idx = i
        break

if sec11_idx is None:
    print("ERROR: Could not find 'Section 11' heading.")
    exit(1)

print(f"Found Section 11 at paragraph {sec11_idx}, next H1 at {next_h1_idx}")

# Remove old Section 11 body
paras_to_remove = paras[sec11_idx + 1 : next_h1_idx]
for p in paras_to_remove:
    p._element.getparent().remove(p._element)

anchor = paras[sec11_idx]._element

# ── Element factories ─────────────────────────────────────────────────────────

def make_h2(text):
    p = doc.add_heading(text, level=2)
    el = p._element; el.getparent().remove(el); return el

def make_h3(text):
    p = doc.add_heading(text, level=3)
    el = p._element; el.getparent().remove(el); return el

def make_body(text):
    p = doc.add_paragraph(text)
    el = p._element; el.getparent().remove(el); return el

def make_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    el = p._element; el.getparent().remove(el); return el

def make_code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text if text else ' ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    el = p._element; el.getparent().remove(el); return el

def make_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x5f, 0xaf)
    el = p._element; el.getparent().remove(el); return el

def make_acdl_placeholder(label):
    p = doc.add_paragraph()
    run = p.add_run(f'[ INSERT ACDL VISUAL: {label} ]')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    run.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)
    el = p._element; el.getparent().remove(el); return el

def make_spacer():
    p = doc.add_paragraph('')
    el = p._element; el.getparent().remove(el); return el

# ── New Section 11 content ────────────────────────────────────────────────────

new_elements = []

# Extension 1
new_elements.append(make_h2('Extension 1 (IMPLEMENTED): Autonomous Agentic Loop (--agentic)'))

new_elements.append(make_h3('Description'))
new_elements.append(make_body(
    'The user gives a high-level goal; the agent plans and executes commands step by step, '
    'observes output, and decides what to do next — without the user specifying each step. '
    'This is a real agent capability: the llm controls both planning and tool use.'
))

new_elements.append(make_h3('Implementation'))
new_elements.append(make_body('A separate system prompt (_AGENTIC_SYSTEM) defines three response types:'))
new_elements.append(make_bullet('"agentic_step" — run a command, observe output, continue'))
new_elements.append(make_bullet('"agentic_done" — goal reached, stop'))
new_elements.append(make_bullet('"agentic_ask" — ask the user a question, then continue'))
new_elements.append(make_body(
    'run_agentic() maintains a growing messages list. After each step, the execution result '
    'is appended and the llm is called again. Max 10 steps.'
))

new_elements.append(make_h3('ACDL — Agentic mode (substep @I within turn @T)'))
for line in [
    'DoitAgentic[@T, @I]: {',
    '    S: {',
    '        AGENTIC_SYSTEM_PROMPT',
    '        env.cwd[@T]',
    '        sys.shell_name',
    '        If sys.memory != empty {',
    '            sys.memory',
    '        }',
    '    }',
    '    U: env.goal[@T]',
    '    If @I > 1 {',
    '        ForEach(@i: range(1, @I-1)) {',
    '            A: resp.agentic_response[@T.i]',
    '            Switch resp.agentic_type[@T.i] {',
    '                Case "agentic_step" {',
    '                    U: env.execution_result[@T.i]',
    '                }',
    '                Case "agentic_ask" {',
    '                    U: env.user_answer[@T.i]',
    '                }',
    '            }',
    '        }',
    '    }',
    '}',
]:
    new_elements.append(make_code(line))
new_elements.append(make_acdl_placeholder('DoitAgentic — substep @I, growing context with execution feedback'))

new_elements.append(make_h3('Agentic system prompt'))
for line in [
    'You are doit running in AGENTIC mode. The user has given you a high-level goal.',
    'Accomplish it step by step by issuing shell commands, observing output, and deciding next steps.',
    '',
    'Respond with ONE JSON object per turn:',
    '  {"type":"agentic_step","command":"<cmd>","description":"<why>","dangerous":<bool>}',
    '  {"type":"agentic_done","summary":"<what was accomplished>"}',
    '  {"type":"agentic_ask","question":"<your question>"}',
    '',
    'Never loop forever: stop with agentic_done after 3 failed attempts.',
    'Current directory: {cwd}  |  Shell: {shell_name}',
]:
    new_elements.append(make_code(line))

new_elements.append(make_h3('Real interaction (groq/llama-3.1-8b-instant)'))
new_elements.append(make_label('doit --agentic "find all python files in this directory and count how many there are"'))
for line in [
    '[Agentic mode - goal: find all python files in this directory and count how many there are]',
    '',
    '-- Agentic step 1/10 --',
    "$ find . -name '*.py' | wc -l",
    '  -> Find all python files in the current directory and count them',
    '1',
    '',
    '-- Agentic step 2/10 --',
    '$ echo $?',
    '  -> Check the return code of the previous command',
    '0',
    '',
    '-- Agentic step 3/10 --',
    '$ echo 1',
    '  -> Print the number of python files found',
    '1',
    '',
    '-- Agentic step 4/10 --',
    '[Done] Found 1 python file in the current directory',
]:
    new_elements.append(make_code(line))
new_elements.append(make_body(
    'The agent ran find, verified the exit code, then confirmed the result before issuing '
    'agentic_done. Steps 2 and 3 were redundant — the llm over-verified — but the final '
    'answer was correct.'
))

new_elements.append(make_h3('Limitations'))
new_elements.append(make_bullet(
    'The llm does not always issue agentic_done when the goal is achieved — in testing, '
    'after completing a "check python version then list files" task the model kept running '
    'extra steps (cd into unrelated directories) instead of stopping. This is a model '
    'quality issue with smaller models like llama-3.1-8b-instant, not a framework bug.'
))
new_elements.append(make_bullet(
    'cd commands have no lasting effect — each command runs in its own subprocess, so '
    'directory changes do not persist between agentic steps.'
))

# Extension 2
new_elements.append(make_h2('Extension 2 (described): Context Compaction'))
new_elements.append(make_body(
    'For long sessions, history grows unboundedly. A compaction step would periodically '
    'summarise older turns into a compact narrative ("user navigated to ~/proj, installed '
    'dependencies, ran tests — all succeeded"), replacing them with a single summary entry. '
    'This directly addresses the 10-entry cap limitation without losing semantic context.'
))

# Extension 3
new_elements.append(make_h2('Extension 3 (described): Project Profiles'))
new_elements.append(make_body(
    'Each directory could have a .doit_profile file (similar to CLAUDE.md) with project-specific '
    'instructions: preferred test command, deployment pipeline, known gotchas. When doit is '
    'invoked, it reads any .doit_profile in the current or parent directories and injects it '
    'as "Project context" into the system message, making the agent behave differently per project.'
))

new_elements.append(make_spacer())

# Insert in forward order
for el in reversed(new_elements):
    anchor.addnext(el)

doc.save(DEST)
print(f"Saved: {DEST}")
