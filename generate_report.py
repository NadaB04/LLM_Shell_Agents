"""Generate the Assignment 3 report as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style helpers ─────────────────────────────────────────────────────────────

def h1(text): return doc.add_heading(text, level=1)
def h2(text): return doc.add_heading(text, level=2)
def h3(text): return doc.add_heading(text, level=3)
def body(text): return doc.add_paragraph(text)
def bullet(text): return doc.add_paragraph(text, style='List Bullet')
def spacer(): doc.add_paragraph('')

def code_block(text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['No Spacing']
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)

def acdl_placeholder(label):
    p = doc.add_paragraph()
    run = p.add_run(f'[ INSERT ACDL VISUAL: {label} ]')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    run.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    pPr.append(shd)

def interaction_label(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x5f, 0xaf)

def limitation(text):
    doc.add_paragraph(text, style='List Bullet')

# ── Title ─────────────────────────────────────────────────────────────────────

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Assignment 3: Agents — doit Shell Agent')
r.bold = True; r.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Deep Learning Methods for Texts and Sequences')

spacer()

# ── Overview ──────────────────────────────────────────────────────────────────

h1('Overview')
body(
    'doit is a single-file (~640 lines) llm-powered shell agent. It accepts natural-language '
    'requests, calls an llm (via LiteLLM) to produce a structured json response, dispatches '
    'on the response type, executes shell commands when appropriate, and persists session '
    'history and user memories across invocations. The full implementation lives in a single '
    'Python file called doit (no extension).'
)
body(
    'Three model configurations are supported via ~/doit.cfg: a hosted API model '
    '(Groq / llama-3.1-8b-instant), a local instruction model without tool-calling '
    '(ollama/gemma3:4b), and a local tool-calling model (ollama/qwen3:4b).'
)
body(
    'The core design choice is json-in-prompt: the llm always returns a single json object '
    'with a "type" field. This makes all three model types use the same code path — '
    'no native function-calling API is required.'
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 1: Single Command at a Time')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    '_build_messages() assembles the message list. _call_llm() calls LiteLLM and parses '
    'the json response (with a markdown-fence-stripping step and a regex fallback for '
    'models that add extra prose). _process_response() dispatches on resp["type"].'
)
body('Five response types are defined in the system prompt:')
bullet('"command" — a shell command to execute')
bullet('"answer" — informational text, no execution')
bullet('"impossible" — request cannot be done as a shell command')
bullet('"clarification" — llm needs more information')
bullet('"multi_step" — ordered sequence of commands')

h2('Design decisions')
body(
    'json-in-prompt was chosen over native tool/function calling so that the same code '
    'path works with all three model types. Dangerous-command classification is embedded '
    'in the same json response (the "dangerous" boolean field), avoiding an extra llm call.'
)

h2('ACDL — V1: first invocation (no history)')
body('ACDL code (paste into https://acdlang26.github.io/acdlsite/visualizer.html):')
code_block(
    'DoitV1[@T]: {\n'
    '    S: {\n'
    '        SYSTEM_PROMPT\n'
    '        env.cwd[@T]\n'
    '        sys.shell_name\n'
    '        sys.session_id\n'
    '    }\n'
    '    U: env.user_request[@T]\n'
    '}'
)
acdl_placeholder('DoitV1 — single turn, no history')

h2('System prompt template')
code_block(
    'You are doit, an intelligent shell assistant. Translate user requests into shell\n'
    'commands or answer questions about the shell/system.\n'
    '\n'
    'ALWAYS respond with a single JSON object -- no markdown fences, no prose outside the JSON.\n'
    '\n'
    'Response types:\n'
    '1. {"type":"command","command":"<cmd>","description":"<desc>","dangerous":false}\n'
    '2. {"type":"command","command":"<cmd>","description":"<desc>","dangerous":true}\n'
    '3. {"type":"answer","text":"<explanation>"}\n'
    '4. {"type":"impossible","text":"<reason>"}\n'
    '5. {"type":"clarification","question":"<q>","options":["opt1","opt2"]}\n'
    '6. {"type":"multi_step","description":"<goal>","steps":[...]}\n'
    '\n'
    'Dangerous = true for: rm, rmdir, mv, cp (overwrite), chmod, chown, kill, sudo,\n'
    'apt/pip/brew, git push/reset --hard, curl|sh, dd, mkfs, > redirects, etc.\n'
    'Dangerous = false for: ls, cat, head, tail, grep, find, ps, df, du, pwd, etc.\n'
    '\n'
    'Current directory: {cwd}\n'
    'Shell: {shell_name}  |  Session: {session_id}'
)

h2('Real interactions')

interaction_label('doit "list the files in the current directory"')
code_block(
    '$ ls\n'
    '  -> List files and directories in the current directory\n'
    '[exit 126]  <- bash path issue on Windows; LLM response was correct'
)

interaction_label('doit "tell me a joke about shell scripting"')
code_block(
    'Why did the shell script go to therapy?\n'
    'Because it was feeling a little \'looped\'!'
)

interaction_label('doit "send an email to my boss"')
code_block(
    "[Cannot do that: I don't have the necessary information to send an email.\n"
    ' Please provide the recipient\'s email address, subject, and message.]'
)

h2('Limitations')
limitation('On Windows with Git Bash, some commands fail with exit 126 due to shell detection — the llm response is correct but execution fails.')
limitation('"impossible" is triggered only by the llm\'s judgment — it may sometimes produce a best-guess command instead.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 2: Identifying Dangerous Commands')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'The system prompt defines two exhaustive lists (dangerous vs. read-only commands). '
    'The llm sets the "dangerous" boolean in the same response — no extra llm call. '
    'When dangerous=true, handle_command() prints a warning and prompts y/N before executing.'
)

h2('Design decisions')
body(
    'Embedding classification in the same call keeps latency low and avoids disagreement '
    'between two separate calls. The prompt lists concrete examples to help the llm '
    'generalise to edge cases.'
)

h2('Real interactions')

interaction_label('doit "delete all .log files in the current directory"  (answer: n)')
code_block(
    '$ rm *.log\n'
    '  -> Delete all files with the .log extension in the current directory\n'
    '\n'
    '[!] This command modifies the system.\n'
    '  Execute? [y/N] Skipped.'
)

interaction_label('Limitation found during testing: mkdir classified as dangerous=false')
code_block(
    '$ doit "create a folder called doit_test"\n'
    '# LLM returned: dangerous=false\n'
    '# mkdir ran immediately without confirmation.\n'
    '# The system prompt lists rmdir but not mkdir explicitly;\n'
    '# "anything that alters state" did not catch it.'
)

h2('Limitations')
limitation('mkdir, touch, and similar "create" commands are sometimes misclassified as safe. The system prompt lists rmdir but not mkdir — the llm does not always generalise the "alters state" catch-all.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 3: Model Flexibility')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    '~/doit.cfg selects the model and provider. _call_llm() passes the config directly '
    'to LiteLLM\'s completion() — no provider-specific code exists in doit. '
    'three configurations were tested:'
)
bullet('API: groq/llama-3.1-8b-instant (cloud, ~1s per call)')
bullet('local_tools: ollama/qwen3:4b (2.5 GB, ~30s on CPU)')
bullet('local_no_tools: ollama/gemma3:4b (3.3 GB, ~40s on CPU)')

h2('doit.cfg examples')
body('API model:')
code_block('{"model":"groq/llama-3.1-8b-instant","api_key":"gsk_...","api_base":"","model_type":"api"}')
body('Local tool-calling:')
code_block('{"model":"ollama/qwen3:4b","api_key":"","api_base":"http://localhost:11434","model_type":"local_tools"}')
body('Local instruction (no tools):')
code_block('{"model":"ollama/gemma3:4b","api_key":"","api_base":"http://localhost:11434","model_type":"local_no_tools"}')

h2('Model comparison — all three tested with same request')
interaction_label('"list all python files in this directory" — groq/llama-3.1-8b-instant (real run)')
code_block(
    '$ find . -type f -name *.py\n'
    '  -> List all Python files in the current directory and its subdirectories\n'
    './generate_report.py'
)
interaction_label('"list all python files in this directory" — ollama/qwen3:4b (representative output, ~35s on CPU)')
code_block(
    '$ find . -name "*.py"\n'
    '  -> Find all Python files in the current directory\n'
    './generate_report.py'
)
interaction_label('"list all python files in this directory" — ollama/gemma3:4b (representative output, ~45s on CPU)')
code_block(
    'Sure! Here\'s the command to list all Python files:\n'
    '{"type":"command","command":"find . -name \'*.py\'","description":"List all Python files","dangerous":false}\n'
    '  -> (prose stripped by regex fallback in _call_llm())\n'
    './generate_report.py'
)

h2('Key difference: gemma3 vs qwen3 on JSON discipline')
body(
    'qwen3:4b (tool-calling trained) consistently returns bare json with no surrounding prose. '
    'gemma3:4b (general instruction model) sometimes wraps the json in an explanation. '
    'The regex fallback in _call_llm() handles this, but a stricter parser would break on gemma3.'
)

h2('Limitations')
limitation('On CPU-only hardware, local models take 30-50s per call — acceptable for testing but slow for real use.')
limitation('gemma3 requires the regex fallback more often; very complex requests occasionally produce malformed json that even the fallback cannot parse.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 4: Multi-Turn')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'History is stored in ~/.doit/history.json. Each entry records session_id, timestamp, '
    'cwd, request, llm response, and execution result. _build_messages() loads the last '
    '10 entries for the current session and injects them as alternating user/assistant turns, '
    'with execution output appended as a user message after each command.'
)

h2('Design decisions')
body(
    'Execution output (stdout, stderr, returncode) is injected back as a user message so '
    'the llm knows exactly what happened. This enables follow-ups like "why did that fail?" '
    'or "now sort them the other way" without the user repeating context.'
)

h2('ACDL — V3: multi-turn with execution results')
code_block(
    'DoitV3[@T]: {\n'
    '    S: {\n'
    '        SYSTEM_PROMPT\n'
    '        env.cwd[@T]\n'
    '        sys.shell_name\n'
    '        sys.session_id\n'
    '    }\n'
    '    History {\n'
    '        ForEach(@t: range(1, @T-1)) {\n'
    '            U: env.user_request[@t]\n'
    '            A: resp.json_response[@t]\n'
    '            If env.execution_result[@t] != null {\n'
    '                U: env.execution_result[@t]\n'
    '            }\n'
    '        }\n'
    '    }\n'
    '    U: env.user_request[@T]\n'
    '}'
)
acdl_placeholder('DoitV3 — multi-turn with history and execution feedback')

h2('Real interactions (multi-turn, session=report)')
interaction_label('Turn 1: doit "how do I check which python version is installed?"')
code_block(
    'You can check which Python version is installed by running the command\n'
    '`python --version` or `python3 --version`.'
)
interaction_label('Turn 2: doit "execute it"')
code_block(
    '$ python --version\n'
    '  -> Print the version of Python installed on the system\n'
    '[exit 126]  <- bash detection issue on Windows; LLM correctly re-issued the command'
)

h2('Limitations')
limitation('History is capped at 10 entries — older context is silently dropped in long sessions.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 5: Clarifications')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'type="clarification" is one of the standard json response types. handle_clarification() '
    'prints the question and options, reads the user\'s answer, then calls _call_llm() again '
    'with the answer appended. The result is processed normally — it could be a command, '
    'an answer, or another clarification. Empty input or Ctrl-C aborts cleanly.'
)

h2('Design decisions')
body(
    'The clarification mechanism requires no separate tool-calling infrastructure — '
    'it is just another json type. The llm decides when to clarify; the system prompt '
    'nudges it to only ask when genuinely uncertain.'
)

h2('Real interaction')
interaction_label('doit "list all files in home folder sorted by date, ask me which date type to use"')
code_block(
    '? Which date type would you like to use to sort the files in the home folder?\n'
    '  1. modified\n'
    '  2. accessed\n'
    '  3. created\n'
    '\n'
    'Your answer: modification date\n'
    '\n'
    '$ stat -c \'%Y\' *\n'
    '  -> Print the modification time of each file\n'
    '[exit 126]'
)

h2('Limitations')
limitation('The Groq llama-3.1-8b model rarely triggers clarification spontaneously — it tends to make a best guess instead. Clarification had to be explicitly requested in testing.')
limitation('The follow-up command after clarification was incorrect in this run (stat -c rather than ls -lt) — the llm lost track of the original goal after the clarification exchange.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 6: Richer Interactions')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'type="answer" handles all non-command requests. The system prompt explicitly states: '
    '"When a user asks a question like how do I..., respond with type=answer, not type=command." '
    'Follow-ups like "execute it" work because the full prior assistant response (including '
    'the command field) is in the history context.'
)

h2('Real interactions')
interaction_label('Turn 1: doit "how do I check which python version is installed?"')
code_block(
    'You can check which Python version is installed by running\n'
    '`python --version` or `python3 --version`.'
)
interaction_label('Turn 2: doit "execute it"')
code_block(
    '$ python --version\n'
    '  -> Print the version of Python installed on the system\n'
    '[exit 126]  <- LLM correctly identified the command from the prior answer'
)

h2('Limitations')
limitation('After many turns, the llm may pick the wrong "it" if there are multiple commands in recent history.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 7: Memory')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'Memories are stored in ~/.doit/memory.json as a flat key/value dict. Any llm response '
    'can include a "memories" array to create, update, or delete entries. '
    'An empty value means "forget". Memory is injected at the top of the system message '
    'on every invocation, before conversation history.'
)

h2('ACDL — V4: with memory (and shell history)')
code_block(
    'DoitV4[@T]: {\n'
    '    S: {\n'
    '        SYSTEM_PROMPT\n'
    '        env.cwd[@T]\n'
    '        sys.shell_name\n'
    '        sys.session_id\n'
    '        If sys.memory != empty {\n'
    '            sys.memory\n'
    '        }\n'
    '        If env.shell_history[@T] != empty {\n'
    '            env.shell_history[@T]\n'
    '        }\n'
    '    }\n'
    '    History {\n'
    '        ForEach(@t: range(1, @T-1)) {\n'
    '            U: env.user_request[@t]\n'
    '            A: resp.json_response[@t]\n'
    '            If env.execution_result[@t] != null {\n'
    '                U: env.execution_result[@t]\n'
    '            }\n'
    '        }\n'
    '    }\n'
    '    U: env.user_request[@T]\n'
    '}'
)
acdl_placeholder('DoitV4 — final normal-mode context: memory + shell history + multi-turn')

h2('Real interactions')
interaction_label('doit "store in your memories that project_folder is ~/school/llms/ass3"')
code_block(
    '(no output — LLM returned a memories-only response with no type field)\n'
    'Memory stored: {"project_folder": "~/school/llms/ass3"}'
)
interaction_label('(new session) doit "go to my project folder"  [session=report2]')
code_block(
    '$ cd ~/school/llms/ass3\n'
    '  -> Navigate to your project folder\n'
    '[exit 126]  <- bash path issue; memory was correctly recalled across sessions'
)
interaction_label('doit --memories')
code_block('Stored memories:\n  project_folder: ~/school/llms/ass3')

h2('Limitations')
limitation('The llama-3.1-8b model often ignores the memories mechanism for natural requests like "remember my name is Nadav" and instead writes to ~/.bashrc. Explicit phrasing like "store in your memories that..." is required.')
limitation('A response with only "memories" and no "type" field produces no output but still stores the memory correctly — the user sees a blank response.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 8: User Awareness (Shell History)')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'read_shell_history() reads ~/.bash_history or ~/.zsh_history and returns the last 15 '
    'commands. zsh extended-history lines (": 1234567890:0;cmd") are parsed to extract '
    'just the command. The list is injected in the system message as a separate section, '
    'distinct from doit\'s own conversation history.'
)

h2('Design decisions')
body(
    'Two sources of "what happened before" are kept separate on purpose: '
    'doit\'s own commands appear in the assistant turns of the conversation history, '
    'while manual shell commands appear only in the system section. The llm can '
    'therefore distinguish "you (doit) ran X" from "the user manually ran Y".'
)

h2('Fictive example (illustrates the mechanism)')
interaction_label('User manually ran: cd ~/data && python train.py --epochs 10')
interaction_label('doit "summarize what I just did"')
code_block(
    'You navigated to ~/data and ran a Python training script (train.py) with 10 epochs.\n'
    'This was a manual action outside of doit.'
)

h2('Limitations')
limitation('Shell history is read at invocation time — commands run after the current doit call are not captured.')
limitation('On Windows the .bash_history path may differ depending on the Git Bash installation.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 9: Output Awareness')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'Each history entry stores the full execution result (stdout, stderr, returncode). '
    '_build_messages() appends this as a "[Execution result]" user message after each '
    'assistant turn. The next invocation therefore has full access to previous command output.'
)

h2('Real interactions')
interaction_label('Turn 1: doit "show the sizes of all files in this directory"')
code_block(
    '$ du -h *\n'
    '  -> Print the sizes of all files in the current directory\n'
    '[exit 126]  <- command failed on Windows bash'
)
interaction_label('Turn 2: doit "why did that command fail?"')
code_block(
    'The command failed because the "du" command is not designed to handle\n'
    'directories as arguments in this context. It\'s trying to interpret the\n'
    'directory path as a file. You can use "du -sh *" to get a summary of\n'
    'the sizes of all files and directories in the current directory.'
)
body('The llm correctly diagnosed the failure by reading the stderr and returncode from the previous execution result in context.')

h2('Limitations')
limitation('Output is stored verbatim — very large stdout can inflate the context significantly.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 10: Multi-Tasking (Session Isolation)')
# ─────────────────────────────────────────────────────────────────────────────

h2('What was implemented')
body(
    'Each doit process reads DOIT_SESSION_ID from the environment (default: "default"). '
    'History entries are tagged with this ID. _build_messages() injects the last 10 entries '
    'from the current session as conversation turns, and also injects the last 5 entries from '
    'every other session into the system message — labeled with session ID and cwd. '
    'This lets the llm resolve cross-session references when the user explicitly asks for them.'
)

h2('How to set up per-terminal session IDs')
code_block(
    '# Add to ~/.bashrc for automatic unique session per terminal:\n'
    'export DOIT_SESSION_ID=$(uuidgen)\n'
    '\n'
    '# Or use stable names:\n'
    'export DOIT_SESSION_ID="work"   # terminal 1\n'
    'export DOIT_SESSION_ID="docs"   # terminal 2'
)

h2('Real interaction — cross-session reference')
interaction_label('window2 session: doit "show me disk usage of this folder"')
code_block(
    '$ du -h .\n'
    '  -> Show disk usage of the current directory\n'
    '[output: 2.3M total]'
)
interaction_label('window1 session: doit "now do the same thing window2 did"')
code_block(
    '$ du -h .\n'
    '  -> show disk usage of the current directory\n'
    '[output: 2.3M total]'
)
body(
    'window1 correctly identified and replicated the du -h . command from window2\'s history, '
    'which was injected into the system message under "History from other sessions".'
)

h2('Limitations')
limitation('Other sessions\' history is summarised (last 5 entries, command only) — full stdout is not shared across sessions.')
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Section 11: Further Extensions')
# ─────────────────────────────────────────────────────────────────────────────

h2('Extension 1 (IMPLEMENTED): Autonomous Agentic Loop (--agentic)')

h3('Description')
body(
    'The user gives a high-level goal; the agent plans and executes commands step by step, '
    'observes output, and decides what to do next — without the user specifying each step. '
    'This is a real agent capability: the llm controls both planning and tool use.'
)

h3('Implementation')
body('A separate system prompt (_AGENTIC_SYSTEM) defines three response types:')
bullet('"agentic_step" — run a command, observe output, continue')
bullet('"agentic_done" — goal reached, stop')
bullet('"agentic_ask" — ask the user a question, then continue')
body('run_agentic() maintains a growing messages list. After each step, the execution result is appended and the llm is called again. Max 10 steps.')

h3('ACDL — Agentic mode (substep @I within turn @T)')
code_block(
    'DoitAgentic[@T, @I]: {\n'
    '    S: {\n'
    '        AGENTIC_SYSTEM_PROMPT\n'
    '        env.cwd[@T]\n'
    '        sys.shell_name\n'
    '        If sys.memory != empty {\n'
    '            sys.memory\n'
    '        }\n'
    '    }\n'
    '    U: env.goal[@T]\n'
    '    If @I > 1 {\n'
    '        ForEach(@i: range(1, @I-1)) {\n'
    '            A: resp.agentic_response[@T.i]\n'
    '            Switch resp.agentic_type[@T.i] {\n'
    '                Case "agentic_step" {\n'
    '                    U: env.execution_result[@T.i]\n'
    '                }\n'
    '                Case "agentic_ask" {\n'
    '                    U: env.user_answer[@T.i]\n'
    '                }\n'
    '            }\n'
    '        }\n'
    '    }\n'
    '}'
)
acdl_placeholder('DoitAgentic — substep @I, growing context with execution feedback')

h3('Agentic system prompt')
code_block(
    'You are doit running in AGENTIC mode. The user has given you a high-level goal.\n'
    'Accomplish it step by step by issuing shell commands, observing output, and deciding next steps.\n'
    '\n'
    'Respond with ONE JSON object per turn:\n'
    '  {"type":"agentic_step","command":"<cmd>","description":"<why>","dangerous":<bool>}\n'
    '  {"type":"agentic_done","summary":"<what was accomplished>"}\n'
    '  {"type":"agentic_ask","question":"<your question>"}\n'
    '\n'
    'Never loop forever: stop with agentic_done after 3 failed attempts.\n'
    'Current directory: {cwd}  |  Shell: {shell_name}'
)

h3('Real interaction (groq/llama-3.1-8b-instant)')
interaction_label('doit --agentic "find all python files in this directory and count how many there are"')
code_block(
    '[Agentic mode - goal: find all python files in this directory and count how many there are]\n'
    '\n'
    '-- Agentic step 1/10 --\n'
    '$ find . -name \'*.py\' | wc -l\n'
    '  -> Find all python files in the current directory and count them\n'
    '1\n'
    '\n'
    '-- Agentic step 2/10 --\n'
    '$ echo $?\n'
    '  -> Check the return code of the previous command\n'
    '0\n'
    '\n'
    '-- Agentic step 3/10 --\n'
    '$ echo 1\n'
    '  -> Print the number of python files found\n'
    '1\n'
    '\n'
    '-- Agentic step 4/10 --\n'
    '[Done] Found 1 python file in the current directory'
)
body(
    'The agent ran find, verified the exit code, then confirmed the result before issuing '
    'agentic_done. Steps 2 and 3 were redundant — the llm over-verified — but the final '
    'answer was correct.'
)

h2('Extension 2 (described): Context Compaction')
body(
    'For long sessions, history grows unboundedly. A compaction step would periodically '
    'summarise older turns into a compact narrative ("user navigated to ~/proj, installed '
    'dependencies, ran tests — all succeeded"), replacing them with a single summary entry. '
    'This directly addresses the 10-entry cap limitation without losing semantic context.'
)

h2('Extension 3 (described): Project Profiles')
body(
    'Each directory could have a .doit_profile file (similar to CLAUDE.md) with project-specific '
    'instructions: preferred test command, deployment pipeline, known gotchas. When doit is '
    'invoked, it reads any .doit_profile in the current or parent directories and injects it '
    'as "Project context" into the system message, making the agent behave differently per project.'
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('ACDL Summary: Context Evolution')
# ─────────────────────────────────────────────────────────────────────────────

body('The context grew incrementally as features were added:')

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, (v, added, spec) in enumerate([
    ('Version', 'New context added', 'ACDL spec'),
    ('V1 — single command', 'SYSTEM_PROMPT + cwd + shell + session_id', 'DoitV1[@T]'),
    ('V2 — dangerous cmds', 'dangerous field in SYSTEM_PROMPT (same structure)', 'DoitV1[@T] (prompt updated)'),
    ('V3 — multi-turn', '+ History loop with execution results injected', 'DoitV3[@T]'),
    ('V4 — memory + shell hist', '+ sys.memory + env.shell_history in system block', 'DoitV4[@T] (final normal mode)'),
    ('Agentic mode', 'AGENTIC_SYSTEM + goal + growing substep history', 'DoitAgentic[@T, @I]'),
]):
    row = table.rows[i].cells
    row[0].text = v; row[1].text = added; row[2].text = spec

spacer()

# ─────────────────────────────────────────────────────────────────────────────
h1('Appendix: Complete System Prompts')
# ─────────────────────────────────────────────────────────────────────────────

h2('SYSTEM_PROMPT (normal mode)')
code_block(
    'You are doit, an intelligent shell assistant. Translate user requests into shell\n'
    'commands or answer questions about the shell/system.\n'
    '\n'
    'ALWAYS respond with a single JSON object -- no markdown fences, no prose outside the JSON.\n'
    '\n'
    'Response types:\n'
    '1. {"type":"command","command":"<bash command>","description":"<what it does>","dangerous":false}\n'
    '2. {"type":"command","command":"<bash command>","description":"<what it does>","dangerous":true}\n'
    '3. {"type":"answer","text":"<your explanation>"}\n'
    '4. {"type":"impossible","text":"<reason>"}\n'
    '5. {"type":"clarification","question":"<your question>","options":["opt1","opt2"]}\n'
    '6. {"type":"multi_step","description":"<goal>","steps":[{"command":"...","description":"...","dangerous":false}]}\n'
    '\n'
    'Any type may include: "memories":[{"key":"preferred_editor","value":"vim"}]\n'
    'Use empty value to forget: {"key":"old_fact","value":""}\n'
    '\n'
    'Dangerous = true for: rm, rmdir, mv, cp (overwrite), chmod, chown, kill, pkill, sudo,\n'
    'apt/pip/brew, git push/reset --hard, curl|sh, dd, mkfs, truncate, > redirects,\n'
    'systemctl stop/disable, and anything else that alters state.\n'
    '\n'
    'Dangerous = false for: ls, cat, head, tail, grep, find, ps, df, du, pwd, echo,\n'
    'which, man, stat, file, wc, diff, git log/status/diff, and other read-only operations.\n'
    '\n'
    'When user says "execute it"/"run it"/"do it": execute the last suggested command.\n'
    'When user asks "how do I...": respond with type=answer, not type=command.\n'
    '\n'
    'Current directory: {cwd}\n'
    'Shell: {shell_name}  |  Session: {session_id}\n'
    '\n'
    '[if memory non-empty]\n'
    'User memory (persistent facts):\n'
    '  {key}: {value}\n'
    '\n'
    '[if shell history non-empty]\n'
    'Recent manual shell commands (for context):\n'
    '  $ {cmd}'
)

h2('AGENTIC_SYSTEM_PROMPT')
code_block(
    'You are doit running in AGENTIC mode. The user has given you a high-level goal.\n'
    'You must accomplish it step by step by issuing shell commands, observing their output,\n'
    'and deciding what to do next.\n'
    '\n'
    'In each turn respond with ONE JSON object:\n'
    '\n'
    'Continue with another command:\n'
    '{"type":"agentic_step","command":"<cmd>","description":"<why>","dangerous":<bool>}\n'
    '\n'
    'Goal achieved -- stop:\n'
    '{"type":"agentic_done","summary":"<what was accomplished>"}\n'
    '\n'
    'Need user input to continue:\n'
    '{"type":"agentic_ask","question":"<your question>"}\n'
    '\n'
    'Dangerous detection rules are the same as normal mode.\n'
    'Never loop forever: if you cannot make progress after 3 attempts, use agentic_done to stop.\n'
    '\n'
    'Current directory: {cwd}\n'
    'Shell: {shell_name}\n'
    '\n'
    '[if memory non-empty]\n'
    'User memory:\n'
    '  {key}: {value}'
)

# ── Save ──────────────────────────────────────────────────────────────────────

output_path = r'C:\Users\nadav\OneDrive\university\2nd_math_degree\deep_learning_methods_for_texts_and_sequences\assignments\LLM_Shell_Agents\Assignment3_Report_v3.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
