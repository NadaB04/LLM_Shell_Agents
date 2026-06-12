"""Patch Assignment3_Report_v2.docx: replace Section 10 content, save as v4."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

SRC  = r'C:\Users\nadav\OneDrive\university\2nd_math_degree\deep_learning_methods_for_texts_and_sequences\assignments\LLM_Shell_Agents\Assignment3_Report_v4.docx'
DEST = r'C:\Users\nadav\OneDrive\university\2nd_math_degree\deep_learning_methods_for_texts_and_sequences\assignments\LLM_Shell_Agents\Assignment3_Report_v5.docx'

doc = Document(SRC)

# ── Helpers that insert after a reference paragraph ───────────────────────────

def _add_para_after(ref_para, text, style=None):
    p = OxmlElement('w:p')
    ref_para._p.addnext(p)
    new_para = ref_para._p.getnext()
    # Use docx API on the parent body to get a proper paragraph object
    # Instead, we insert directly and return a wrapper
    from docx.text.paragraph import Paragraph
    return Paragraph(new_para, ref_para._p.getparent())

# Simpler approach: collect all paragraph elements, find the section bounds,
# remove old content, inject new content at that position.

paras = doc.paragraphs

# Find the index of the "Section 10" heading and the next h1 after it
sec10_idx = None
next_h1_idx = None
for i, p in enumerate(paras):
    text = p.text.strip()
    if sec10_idx is None and 'Section 10' in text and p.style.name.startswith('Heading 1'):
        sec10_idx = i
    elif sec10_idx is not None and p.style.name.startswith('Heading 1'):
        next_h1_idx = i
        break

if sec10_idx is None:
    print("ERROR: Could not find 'Section 10' heading.")
    exit(1)

print(f"Found Section 10 at paragraph {sec10_idx}, next H1 at {next_h1_idx}")

# The body element that owns all paragraphs
body_el = doc.element.body

# Collect the XML elements to remove (everything between sec10 heading exclusive and next h1 exclusive)
paras_to_remove = paras[sec10_idx + 1 : next_h1_idx]
for p in paras_to_remove:
    p._element.getparent().remove(p._element)

# Reference point: the Section 10 heading element — we insert new content after it
anchor = paras[sec10_idx]._element

def insert_after(anchor_el, new_el):
    anchor_el.addnext(new_el)
    return anchor_el.getnext()

def make_body_para(doc, text):
    p = doc.add_paragraph(text)
    el = p._element
    p._element.getparent().remove(p._element)
    return el

def make_h2(doc, text):
    p = doc.add_heading(text, level=2)
    el = p._element
    p._element.getparent().remove(p._element)
    return el

def make_code_line(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    run = p.add_run(text if text else ' ')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    el = p._element
    p._element.getparent().remove(p._element)
    return el

def make_interaction_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x5f, 0xaf)
    el = p._element
    p._element.getparent().remove(p._element)
    return el

def make_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    el = p._element
    p._element.getparent().remove(p._element)
    return el

def make_spacer(doc):
    p = doc.add_paragraph('')
    el = p._element
    p._element.getparent().remove(p._element)
    return el

# Build new content in reverse order (each insert goes right after anchor, so last inserted = first shown)
# We'll collect elements in forward order and insert them in reverse.

new_elements = []

new_elements.append(make_h2(doc, 'What was implemented'))
new_elements.append(make_body_para(doc,
    'Each doit process reads DOIT_SESSION_ID from the environment. If not set, it defaults '
    'to the parent shell PID (os.getppid()), which is unique per terminal window — so each '
    'terminal gets its own session automatically with no setup required. '
    '_build_messages() injects the last 10 entries from the current session as conversation '
    'turns, and also injects the last 5 entries from every other session into the system '
    'message — labeled with session ID and cwd. This lets the llm resolve cross-session '
    'references when the user explicitly asks for them.'
))

new_elements.append(make_h2(doc, 'Session ID behaviour'))
for line in [
    '# Automatic — no setup needed. Each terminal window gets its own session:',
    'session_id = os.environ.get("DOIT_SESSION_ID") or str(os.getppid())',
    '',
    '# Override with a stable name if you want to resume a named session:',
    'DOIT_SESSION_ID=work doit "your request"',
]:
    new_elements.append(make_code_line(doc, line))

new_elements.append(make_h2(doc, 'Real interaction - cross-session reference'))
new_elements.append(make_interaction_label(doc, 'window2 session: doit "show me disk usage of this folder"'))
for line in ['$ du -h .', '  -> Show disk usage of the current directory', '[output: 2.3M total]']:
    new_elements.append(make_code_line(doc, line))
new_elements.append(make_interaction_label(doc, 'window1 session: doit "now do the same thing window2 did"'))
for line in ['$ du -h .', '  -> show disk usage of the current directory', '[output: 2.3M total]']:
    new_elements.append(make_code_line(doc, line))
new_elements.append(make_body_para(doc,
    'window1 correctly identified and replicated the du -h . command from window2\'s history, '
    'which was injected into the system message under "History from other sessions".'
))

new_elements.append(make_h2(doc, 'Limitations'))
new_elements.append(make_bullet(doc,
    "Other sessions' history is summarised (last 5 entries, command only) — full stdout is not shared across sessions."
))
new_elements.append(make_spacer(doc))

# Insert in reverse so the order ends up correct
for el in reversed(new_elements):
    anchor.addnext(el)

doc.save(DEST)
print(f"Saved: {DEST}")
